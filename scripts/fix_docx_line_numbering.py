# -*- coding: utf-8 -*-
"""Post-process manuscript_main.docx: double line spacing + continuous line numbering (BMC LHD requirement)."""
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\NHANES\manuscript\final\manuscript_main.docx"

doc = docx.Document(SRC)

# 1. Normal style: double spacing for all paragraphs using Normal (pandoc BodyText maps to Normal in default template? set both)
for style_name in ("Normal", "BodyText"):
    try:
        st = doc.styles[style_name]
        st.paragraph_format.line_spacing = 2.0
    except KeyError:
        pass

# 2. Walk all paragraphs: set double spacing explicitly
for p in doc.paragraphs:
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

# 3. Line numbering in section properties (continuous)
sect = doc.sections[0]
sectPr = sect._sectPr
# remove existing lnNumType if any
for el in sectPr.findall(qn('w:lnNumType')):
    sectPr.remove(el)
ln = OxmlElement('w:lnNumType')
ln.set(qn('w:countBy'), '1')
ln.set(qn('w:restart'), 'continuous')
sectPr.append(ln)

# 4. Heading styles: force black (pandoc default template colors headings 0F4761; BMC requires black text)
for style_name in ("Heading 1", "Heading 2", "Heading 3"):
    try:
        st = doc.styles[style_name]
        st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    except KeyError:
        pass

# 5. Page numbering in footer (BMC LHD requirement: "Include line and page numbering")
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field via fldSimple
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')  # 10pt
    rPr.append(sz)
    r.append(rPr)
    t_el = OxmlElement('w:t'); t_el.text = '1'
    r.append(t_el)
    fld.append(r)
    p._p.append(fld)

for section in doc.sections:
    add_page_number(section)

doc.save(r"D:\NHANES\manuscript\final\manuscript_main_lhd.docx")
print("saved manuscript_main_lhd.docx")