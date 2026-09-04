# -*- coding: utf-8 -*-
import re
import sys
from docx import Document

path = sys.argv[1] if len(sys.argv) > 1 else r"D:\NHANES\manuscript\final\BHV15049_WTI_NHANES&CHARLS.docx"
doc = Document(path)
txt = "\n".join(p.text for p in doc.paragraphs)
print("file:", path)

refs = re.findall(r"^(\d+)\.", txt, re.M)
print("reference entries:", len([x for x in refs if x.isdigit()]), "(expect 48)")

# 完整一致版：这些必须存在（与投稿版一致，官方 Q3/Q5 口径）
expect_present = ["Declarations", "Ethics approval and consent", "Competing interests",
                  "no specific grant", "Authors\u2019 contributions", "Use of AI",
                  "zhao_chenguang@outlook.com", "ORCID", "Xing-kai Liu", "Fourth Military Medical University",
                  "10,302", "531 strokes", "9,870", "OR 1.17", "HR 1.12", "90.5", "0.689",
                  "WTSAF2YR", "STROBE", "Trial registration", "IRB00001052"]
for a in expect_present:
    print(("OK     " if a in txt else "MISS!! ") + a)

# lean 版才有机会缺席；完整版里不应缺的章节名
for sec in ["Background", "Methods", "Results", "Discussion", "Conclusions", "List of abbreviations", "References", "Figure legends", "Tables"]:
    print(("sec-OK " if sec in txt else "sec-MISS ") + sec)

print("tables:", len(doc.tables), "(expect 5)")
for t in doc.tables:
    print("  table first row:", [c.text[:18] for c in t.rows[0].cells][:5])
print("chars:", sum(len(p.text) for p in doc.paragraphs))
