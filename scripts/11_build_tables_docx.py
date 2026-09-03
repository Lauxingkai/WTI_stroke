# ============================================================================
# 11_build_tables_docx.py  (Task A prep: journal-ready three-line tables DOCX)
# Builds Table 1/2/3/4/5 as a submission-ready DOCX using python-docx,
# from the audited CSVs (no hand-typed numbers).
# Output: manuscript/final/tables_submission.docx
# Date: 2026-08-16; revised 2026-08-20 (Table 1 added; M1-maximal-sample
# footnotes; Lag-2 sample note)
# ============================================================================
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

RES = r"D:\NHANES\results"
OUT = r"D:\NHANES\manuscript\final\tables_submission.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(9)

def cell_text(v):
    # BMC table rule: no comma separators in numerical values (journal guideline:
    # "Commas should not be used to indicate numerical values").
    if isinstance(v, float) and v == int(v):
        return str(int(v))
    if isinstance(v, (int,)) and not isinstance(v, bool):
        return str(v)
    return str(v)

def add_table(title, df, footnote):
    h = doc.add_paragraph()
    r = h.add_run(title); r.bold = True; r.font.size = Pt(9)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(c)
        for p in cell.paragraphs:
            p.runs[0].bold = True; p.runs[0].font.size = Pt(8)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = cell_text(v)
            for p in cells[j].paragraphs:
                p.runs[0].font.size = Pt(8)
    f = doc.add_paragraph()
    fr = f.add_run(footnote); fr.font.size = Pt(8)
    doc.add_paragraph()

# ---- Table 1: baseline characteristics by stroke status, two cohorts ----
t1n = pd.read_csv(RES + r"\05f_table1_nhanes_outcome.csv")
t1c = pd.read_csv(RES + r"\05f_table1_charls_outcome.csv")

def wide1(df, cohort):
    out = {}
    for _, r in df.iterrows():
        key = r["variable"]
        grp = r["group"]
        if r["level"] == "median (IQR)":
            out[(key, grp)] = {"level": r["level"], "est": r["est"], "p": r["p"]}
        else:
            block = out.setdefault((key, grp), {})
            block[r["level"]] = r["est"]
            block["p"] = r["p"]
    return out

wN, wC = wide1(t1n, "NHANES"), wide1(t1c, "CHARLS")

def pfmt(p):
    # BMC/LHD convention: P values to 3 decimals (journal family practice, cf. published
    # articles in Lipids in Health and Disease); <0.001 reported as "<0.001".
    p = float(p)
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"

rows = []
vars_order = ["Age, years", "Male", "Education", "BMI, kg/m2", "WTI, cm-mmol/L",
              "Waist circumference, cm", "Triglycerides, mmol/L",
              "Physical activity, min/day", "MVPA, days/week",
              "Ever smoker", "Alcohol drinker", "Hypertension", "Diabetes",
              "Statin use", "Lipid-lowering use", "Antihypertensive use"]
all_vars = {k[0] for k in wN} | {k[0] for k in wC}
seen_vars = set()
for var in vars_order + sorted(all_vars - set(vars_order)):
    if var not in all_vars or var in seen_vars:
        continue
    seen_vars.add(var)
    entries = [wN.get((var, g), {}) for g in ("No stroke", "Stroke")] + \
              [wC.get((var, g), {}) for g in ("No stroke", "Stroke")]
    start = len(rows)
    if var == "Male":
        rows.append(["Male sex",
                     entries[0].get("Male", ""), entries[1].get("Male", ""),
                     pfmt(entries[0].get("p")) if entries[0] else "",
                     entries[2].get("Male", ""), entries[3].get("Male", ""),
                     pfmt(entries[2].get("p")) if entries[2] else ""])
        continue
    if "est" in entries[0] or "est" in entries[2]:
        rows.append([var,
                     entries[0].get("est", ""), entries[1].get("est", ""),
                     pfmt(entries[0]["p"]) if "p" in entries[0] else "",
                     entries[2].get("est", ""), entries[3].get("est", ""),
                     pfmt(entries[2]["p"]) if "p" in entries[2] else ""])
    else:
        srcN = entries[0] if entries[0] else entries[1]
        srcC = entries[2] if entries[2] else entries[3]
        levels = [l for l in srcN if l != "p"] + \
                 [l for l in srcC if l != "p" and l not in srcN]
        for i, lv in enumerate(levels):
            label = f"{var}: {lv}" if i == 0 else f"  {lv}"
            rows.append([label, entries[0].get(lv, ""), entries[1].get(lv, ""), "",
                         entries[2].get(lv, ""), entries[3].get(lv, ""), ""])
    if len(rows) > start:
        # P on the variable's first line
        pN = entries[0].get("p") or entries[1].get("p")
        pC = entries[2].get("p") or entries[3].get("p")
        if "est" not in entries[0] and "est" not in entries[2]:
            rows[start][3] = pfmt(pN) if pN else ""
            rows[start][6] = pfmt(pC) if pC else ""

t1 = pd.DataFrame(rows, columns=["Variable", "NHANES no stroke", "NHANES stroke", "P",
                                 "CHARLS no stroke", "CHARLS stroke", "P"])
add_table(
    "Table 1. Baseline characteristics by stroke status, NHANES 2005-2018 fasting subsample "
    "and CHARLS 2011 cross-sectional wave",
    t1,
    "Continuous variables: median (IQR); categorical: n (weighted %). Design-corrected tests "
    "(Rao-Scott chi-square or weighted Wald). NHANES n = 10,302 (no stroke 9,771; stroke 531), "
    "pooled fasting-subsample weights. CHARLS n = 9,856 minimally adjusted (M1) sample "
    "(no stroke 9,636; stroke 220), 2011 blood weights, community cluster. "
    "CHARLS MVPA in days/week; NHANES physical activity in min/day.")

# ---- Table 2: main associations ----
t2 = pd.read_csv(RES + r"\Table2_main_models.csv")
add_table("Table 2. Main associations of WTI with stroke (per 1-SD), by cohort and model",
          t2, "OR: survey-weighted logistic (quasibinomial), M1 = age/sex, M2 = +race(NHANES)/education/smoking/drinking/BMI, "
              "M3 = +hypertension/diabetes/lipid-lowering/antihypertensive medication/physical activity. "
              "HR: weighted Cox (cluster = community); sHR: Fine-Gray subdistribution (unweighted, iid SE; "
              "community-block bootstrap CIs in Table S3). "
              "n/events are model-fitted counts: CHARLS cross-sectional M1 n = 9,856 (220 events), adjusted models n = 9,206 (178); "
              "CHARLS prospective M1 n = 9,636 (569 events, 131 deaths), adjusted models n = 9,028 (515, 107); "
              "NHANES n = 10,302 (531 events).")

# ---- Table 3: discrimination ----
t3 = pd.read_csv(RES + r"\Table3_discrimination.csv")
add_table("Table 3. Discrimination of seven indices for stroke (base model: age + sex)",
          t3, "AUC: unweighted complete-case logistic predictions, DeLong CI; vs-base column: DeLong test "
              "of each index against the base model (age + sex). NRI/IDI: continuous, bootstrap percentile CIs (B = 1000). "
              "Exploratory; unweighted-calculation caveat in Limitations.")

# ---- Table 4: sensitivity/mediation ----
t4 = pd.read_csv(RES + r"\Table4_sensitivity.csv")
add_table("Table 4. Sensitivity and mediation analyses",
          t4, "E-value per VanderWeele & Ding (rare-outcome approximation); Lag-2 landmark excludes events/deaths within 2 y "
              "(M1 n = 9,545, 539 strokes; M3 n = 8,965, 492 strokes); interval-censored discrete-time person-period model (3 intervals); "
              "mediation: product-of-coefficients, individual-resampling bootstrap (B = 1000), 2011 WTI -> 2015 mediator -> 2018 stroke. "
              "Regression calibration corrected HR from repeated TG pairs (2011-2015, n = 7,482; calibration slope 0.52).")

# ---- Table 5: CHARLS 2015 replication + NHANES NDI prospective layers ----
r15 = pd.read_csv(RES + r"\13_2015_main_models.csv")
rn  = pd.read_csv(RES + r"\13g_ndi_cox_models.csv")

def fmt(row):
    return f"{row.est:.3f} ({row.lo:.3f}-{row.hi:.3f})"

rows5 = []
def row5(layer, model, key, src="r15"):
    r = r15 if src == "r15" else rn
    s = r[(r.layer == key) & (r.model == model)].iloc[0]
    rows5.append([layer, model, fmt(s), pfmt(s.p), f"{int(s.n)}", f"{int(s.events)}"])

row5("CHARLS 2015 cross", "CM1", "cross")
row5("CHARLS 2015 cross", "CM2", "cross")
row5("CHARLS 2015 cross", "CM3", "cross")
row5("CHARLS 2015 cross, alt weight", "CA3", "cross-altw")
row5("CHARLS 2015 cross, physician-confirmed", "CP1", "cross-phys")
row5("NHANES NDI all-cause", "AM1", "all-cause", src="rn")
row5("NHANES NDI all-cause", "AM3", "all-cause", src="rn")
row5("NHANES NDI stroke death", "SM1", "stroke-death", src="rn")
row5("NHANES NDI stroke death", "SM3", "stroke-death", src="rn")
row5("NHANES NDI stroke death, Fine-Gray", "M3", "stroke-death-FG", src="rn")

t5 = pd.DataFrame(rows5, columns=["Layer", "Model", "OR/HR/sHR (95% CI)", "P", "n", "events"])
add_table("Table 5. Replication (CHARLS 2015) and NHANES NDI mortality associations of WTI with stroke (per 1-SD)",
          t5, "CHARLS 2015: survey-weighted logistic (quasibinomial), design = community cluster + urban/rural strata, "
              "weight = 2015 blood weight normalized; M1 age/sex, M2 + education/smoking/drinking/BMI, "
              "M3 + hypertension/diabetes/lipid-lowering/antihypertensive medication/physical activity. "
              "Physician-confirmed stroke: doctor-told verification records (33 events). "
              "NHANES NDI: survey-weighted cause-specific Cox (cluster = cycle-specific PSU, strata likewise, pooled weights), "
              "follow-up through Dec 31, 2019 (median 6.9 y; 74,744 person-years); stroke death = underlying cause I60-I69. "
              "Events are model-fitted (complete cases); cohort totals: all-cause 1,428 and cerebrovascular 83 deaths (n = 10,289); "
              "stroke-death M1 fitted on 83 events, M3 on 82. "
              "Fine-Gray: unweighted, other deaths as competing events. Full row sets in Supplementary Tables S4-S5.")

doc.save(OUT)
print("DONE ->", OUT)
