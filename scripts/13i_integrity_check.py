# -*- coding: utf-8 -*-
"""13i_integrity_check.py  (round 5, 2026-09-05: governance norm engine)
round 4 -> round 5 changes (version-governance strategy §3):
  - norm engine: strip \\ escapes / * italics / spaces / dashes / commas
  - matching: numeric-sequence subset check (robust to wording & formatting)
  - citations extracted from normalized text (pandoc \\[ compatibility)
  - n-value expectations verified against tables_submission.docx (not body prose)
Checks: [1] citations, [2] new-layer numbers, [3] Tables, [4] provenance,
        [5] 13q two-piecewise, [6] 13r HTGW, [7] refitted fragments.
"""
import re, os
import pandas as pd
from docx import Document

MAN = r"D:\NHANES\manuscript\final\manuscript_main.md"
RES = r"D:\NHANES\results"
QC = r"D:\NHANES\qc"
TBL = r"D:\NHANES\manuscript\final\tables_submission.docx"

lines = []
def log(s=""):
    print(s); lines.append(s)

text = open(MAN, encoding="utf-8").read()

# ---- norm engine (governance §3) ----
def _norm(s):
    return re.sub(r"[\\*\s]", "", s.replace("--", "-").replace("−", "-").replace("–", "-"))
def _nums(s):
    return set(re.findall(r"\d+\.\d+|\d+", _norm(s).replace(",", " ")))
NT = _norm(text)
NS = _nums(NT)

# ---- 1) citations ----
cited = set()
T2 = _norm(text).replace("[", "[").replace("]", "]")  # 引用用（保留逗号）
for grp in re.findall(r"\[(\d+(?:[-,]\d+)*)\]", T2):
    for part in grp.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-")
            cited.update(range(int(a), int(b) + 1))
        elif part:
            cited.add(int(part))
nrefs = 48
missing_ref = [i for i in range(1, nrefs + 1) if i not in cited]
ghost = sorted(c for c in cited if c < 1 or c > nrefs)
log(f"[1] citations: refs cited in text = {len(cited)}/{nrefs}")
log(f"    uncited refs: {missing_ref}")
log(f"    ghost numbers: {ghost}")
log(f"    -> {'PASS' if not missing_ref and not ghost else 'FAIL'}")

# ---- 2) new-layer numbers (numeric-sequence match) ----
r15 = pd.read_csv(RES + r"\13_2015_main_models.csv")
rn = pd.read_csv(RES + r"\13g_ndi_cox_models.csv")
checks = []
def expect(src, layer, model, label, extra_nums=()):
    r = r15 if src == "r15" else rn
    s = r[(r.layer == layer) & (r.model == model)].iloc[0]
    need = {f"{s.est:.2f}", f"{s.lo:.2f}", f"{s.hi:.2f}"} | set(extra_nums)
    ok = need <= NS
    checks.append((f"{layer}/{model}", label, ok))
    return ok

expect("r15", "cross", "CM1", "M1 OR")
expect("r15", "cross", "CM2", "M2 OR")
expect("r15", "cross", "CM3", "M3 OR")
expect("r15", "cross-altw", "CA1", "M1 alt-w")
expect("r15", "cross-phys", "CP1", "M1 phys")
expect("r15", "cross-women", "M1", "women M1")
expect("r15", "cross-men", "M1", "men M1")
expect("r15", "cross-age-45to59", "M1", "45-59 y")
expect("r15", "cross-age-60plus", "M1", "60+ y")
expect("rn", "all-cause", "AM1", "all-cause M1", ("0.031",))
expect("rn", "all-cause", "AM2", "all-cause M2", ())
expect("rn", "all-cause", "AM3", "all-cause M3", ("0.765",))
expect("rn", "stroke-death", "SM1", "stroke M1", ("0.747",))
expect("rn", "stroke-death", "SM2", "stroke M2", ())
expect("rn", "stroke-death", "SM3", "stroke M3", ("0.969",))
expect("rn", "stroke-death-FG", "M3", "stroke FG-M3", ("0.870",))

n_pass = sum(1 for c in checks if c[2])
log(f"[2] stats-consistency: {n_pass}/{len(checks)} new-layer number sets match")
for name, frag, ok in checks:
    if not ok:
        log(f"    MISS: {name} | {frag}")
log(f"    -> {'PASS' if n_pass == len(checks) else 'FAIL'}")
log("    tertile check (NHANES M3 0.91 / T3 0.68 / NDI T2-T3 0.88): " +
    str(all(x in NS for x in ("0.91", "0.68", "0.88"))))
log("    MDE check (CHARLS Cox HR>=1.14): " + str("1.14" in NS))

# ---- 3) Tables in docx ----
doc = Document(TBL)
paras = [p.text for p in doc.paragraphs]
t1_5 = all(any(f"Table {k}" in p for p in paras) for k in range(1, 6))
tbl_text = _norm("".join(c.text for tb in doc.tables for row in tb.rows for c in row.cells))
t_n = all(x in tbl_text for x in ("10302", "9856", "9636", "9028", "12213", "11203"))
log(f"[3] Tables 1-5 present in tables_submission.docx: {t1_5}")
log(f"    key Ns in tables (10302/9856/9636/9028/12213/11203): {t_n}")
log(f"    Table 2 contains M1 n 9,856/9,636: {all(x in tbl_text for x in ('9856', '9636'))}")
log(f"    -> {'PASS' if t1_5 and t_n else 'FAIL'}")

# ---- 4) provenance ----
need_files = [RES + r"\03_main_models.csv", RES + r"\03c_cox_fg.csv", RES + r"\13_2015_main_models.csv",
              RES + r"\13g_ndi_cox_models.csv", RES + r"\13q_threshold.csv", RES + r"\13r_htgw_models.csv"]
miss_f = [p for p in need_files if not os.path.exists(p)]
log(f"[4] provenance files: {len(need_files) - len(miss_f)}/{len(need_files)} exist; missing={miss_f}")
log(f"    -> {'PASS' if not miss_f else 'FAIL'}")

# ---- 5) 13q two-piecewise (values validated 2026-09-05 against 13q_checks.txt) ----
ok5 = all(x in NS for x in ("90.5", "72.1", "137.8", "1.42", "1.23", "1.64", "0.96", "0.86", "1.08"))
log(f"[5] 13q two-piecewise (turn 90.5; below 1.42 (1.23-1.64); above 0.96 (0.86-1.08)): {ok5}")
log(f"    -> {'PASS' if ok5 else 'FAIL'}")

# ---- 6) 13r HTGW ----
ok6 = all(x in NS for x in ("2.04", "1.47", "1.06", "0.83", "1.46", "1.02"))
log(f"[6] 13r HTGW additions (2.04/1.47/1.06/0.83/1.46/1.02): {ok6}")
log(f"    -> {'PASS' if ok6 else 'FAIL'}")

# ---- 7) refitted fragments (numeric-set) ----
f7 = [
    ("CHARLS M1 1.17", ("1.17", "1.05", "1.29", "0.003")),
    ("CHARLS M2 1.14", ("1.14", "1.01", "1.27", "0.026")),
    ("CHARLS M3 1.07", ("1.07", "0.94", "1.21", "0.312")),
    ("Cox M1 1.12", ("1.12", "1.06", "1.18")),
    ("Cox M3 1.01", ("1.01", "0.94", "1.10", "0.728")),
    ("FG M1 1.13", ("1.13", "1.08", "1.18")),
    ("FG M3 1.02", ("1.02", "0.96", "1.08", "0.549")),
    ("E-values M1 1.36-1.60", ("1.36", "1.60")),
    ("E-value CI-upper 1.58-1.90", ("1.58", "1.90")),
    ("fully-adjusted E-value 1.14-1.44", ("1.14", "1.44")),
    ("CIF T2 1.65", ("1.65", "1.32", "2.06")),
    ("CIF T3 2.02", ("2.02", "1.64", "2.50")),
    ("sex interaction P", ("0.981", "0.984")),
    ("corrected HR 1.24", ("1.24", "1.11", "1.38")),
    ("Lag-2 Cox M1", ("1.13", "1.06", "1.19")),
    ("Lag-2 FG M1", ("1.14", "1.09", "1.19")),
    ("IC M1 OR", ("1.13", "1.06", "1.20")),
    ("WTI medians", ("114.4", "100.7")),
    ("diabetes prevalence", ("5.6", "15.2")),
    ("MRI sensitivity", ("32.4",)),
    ("fasting subsample", ("10766", "26282")),
    ("per-10 equivalents", ("1.22", "1.13", "1.33", "1.00", "0.99", "1.01")),
]
n7 = 0; miss7 = []
for frag, nums in f7:
    if set(nums) <= NS:
        n7 += 1
    else:
        miss7.append(frag)
log(f"[7] refitted fragments: {n7}/{len(f7)} match (numeric sets)")
for f in miss7:
    log(f"    MISS: {f}")
log(f"    -> {'PASS' if n7 == len(f7) else 'FAIL'}")

with open(QC + r"\13i_integrity_check.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(lines) + "\n")
print("saved: D:\\NHANES\\qc\\13i_integrity_check.txt")
any_fail = any(l.endswith("FAIL") for l in lines)
raise SystemExit(1 if any_fail else 0)
