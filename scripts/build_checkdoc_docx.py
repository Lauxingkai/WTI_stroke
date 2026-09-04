# -*- coding: utf-8 -*-
"""
Build a plagiarism-check-ready manuscript DOCX from the final submission files.
Default (recommended, per official iThenticate guidance): the file must match
the version sent to the journal EXACTLY - full text, funding statement,
acknowledgements, references, tables all included
(ithenticate.topeditsci.com/website/other/faq?type=check, Q3/Q5).
Optional --lean: drop author block and Declarations template sentences for a
cleaner self-check number.
"""
import re
import subprocess
import pathlib
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

BASE = pathlib.Path(r"D:\NHANES\manuscript\final")
SRC = BASE / "manuscript_main.md"
TBL_DOCX = BASE / "tables_submission.docx"
OUT = BASE / "BHV15049_WTI_NHANES&CHARLS.docx"
TMP_MD = BASE / ".checkdoc_tmp.md"


def split_sections(text: str) -> dict:
    lines = text.splitlines()
    title = lines[0]
    hr_idx = next(i for i, ln in enumerate(lines) if ln.startswith("---"))
    body_start = lines[hr_idx:]
    d_idx = next(i for i, ln in enumerate(body_start) if ln.startswith("## Declarations"))
    r_idx = next(i for i, ln in enumerate(body_start) if ln.startswith("## References"))
    f_idx = next(i for i, ln in enumerate(body_start) if ln.startswith("## Figure legends"))
    return {
        "title": title,
        "author_block": "\n".join(lines[1:hr_idx]),
        "body": "\n".join(body_start[:d_idx]),
        "declarations": "\n".join(body_start[d_idx:r_idx]),
        "references": "\n".join(body_start[r_idx:f_idx]),
        "legends": "\n".join(body_start[f_idx:]),
    }


def extract_tables_md() -> str:
    res = subprocess.run(
        ["pandoc", str(TBL_DOCX), "-t", "markdown", "--wrap=none"],
        capture_output=True, text=True, encoding="utf-8",
    )
    if res.returncode != 0:
        raise RuntimeError(f"pandoc tables failed: {res.stderr[:500]}")
    return res.stdout


def main() -> None:
    args = sys.argv[1:]
    lean = "--lean" in args
    out_file = OUT
    if "--out" in args:
        out_file = pathlib.Path(args[args.index("--out") + 1])
    parts = split_sections(SRC.read_text(encoding="utf-8"))
    tables_md = extract_tables_md()
    # escape leading numbers of reference lines so pandoc renders plain
    # paragraphs instead of auto-numbered list items (parsers may miss the
    # numbers in auto-numbered Word lists)
    refs = re.sub(r"^(\d+)\.", r"\1\\.", parts["references"].strip(), flags=re.M)
    ref_lines = [ln.rstrip() for ln in refs.splitlines() if ln.strip()]
    refs = "\n\n".join(ref_lines)
    chunks = [parts["title"]]
    if not lean:
        chunks.append(parts["author_block"].strip())
    chunks.append(parts["body"].rstrip())
    if not lean:
        chunks.append(parts["declarations"].strip())
    chunks.append(refs)
    chunks.append(parts["legends"].strip())
    chunks.append("# Tables\n\n" + tables_md.strip())
    md = "\n\n".join(chunks) + "\n"
    TMP_MD.write_text(md, encoding="utf-8")
    subprocess.run(["pandoc", str(TMP_MD), "-o", str(out_file), "--wrap=none"], check=True)

    doc = Document(str(out_file))
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(6)
    for sname in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            st = doc.styles[sname]
        except KeyError:
            continue
        st.font.name = "Times New Roman"
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            for r in p.runs:
                r.font.name = "Times New Roman"
    doc.save(str(out_file))

    chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"OK -> {out_file}")
    print(f"paragraphs={len(doc.paragraphs)} chars={chars} tables={len(doc.tables)}")
    heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    print("headings:", heads)
    alltxt = " ".join(p.text for p in doc.paragraphs)
    print("author info leaked:", any(k in alltxt for k in ("Xing-kai", "edelweiss", "ORCID")))
    print("Declarations present:", "Declarations" in alltxt)


if __name__ == "__main__":
    main()
